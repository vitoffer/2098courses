import os
from shutil import copyfileobj
from typing import Annotated

# import openpyxl as xl
from docx.api import Document
from docx.table import _Row
from sqlmodel import Session, select
from models.models import Course, Schedule
from models.database import get_session
from .models import DeleteTables, UploadTabelToDb, CourseType, ScheduleType
from conf.settings import BASE_DIR
from fastapi import APIRouter, UploadFile, Depends, HTTPException, status
from routers.auth import get_current_active_user, User
from conf.depencies import super_user

import bs4
import requests

TABLES_DIR = BASE_DIR / 'cache'
SCHOOL = 'https://sch2098s.mskobr.ru/o-nas/pedagogicheskii-sostav'

KEYS = [
    'teacher',
    'name',
    'for_ages',
    'class',
    'address',
    'is_paid',
    'description',
]

DAYS = [
    'monday',
    'tuesday',
    'wednesday',
    'thursday',
    'friday',
    'saturday',
]

router = APIRouter(prefix='/admin', tags=['admin'])


def is_super_user(user: User):
    if not user.is_super_user:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail='Access denied',
            headers={'WWW-Authenticate': 'Bearer'},
        )
    return True


def convert_to_models(args):
    k, row, address, is_paid = args
    if k <= 1:
        return None, None
    text = [cell.text for cell in row.cells[:4]]
    days = [cell.text for cell in row.cells[4:]]
    text.extend((address, is_paid, ''))
    return {KEYS[_]: text[_] for _ in range(len(KEYS))}, {
        DAYS[_]: days[_] for _ in range(len(DAYS)) if days[_] != ''
    }


# Парсер для получения списка учителей
# @router.get('/get-teacher-list')
# async def get_tachers():
#     res = requests.get(SCHOOL)
#     with open(f'{TABLES_DIR}/index.txt', 'wb') as f:
#         f.write(res.content)
#     with open(f'{TABLES_DIR}/index.txt', 'r') as f:
#         s = f.read()
#     site = bs4.BeautifulSoup(s, 'html.parser')
#     divs = site.findAll('div')
#     # teachers = bs4.Tag()
#     for i in divs:
#         try:
#             if i['class'] == ['container-sortable2', 'group-teachers'] :
#                 teachers = i
#                 break
#         except Exception:
#             print('')
#     divs = bs4.BeautifulSoup(str(teachers.prettify()), 'html.parser')
#     return {}


@router.post('/upload_table')
async def upload_table(
    table: UploadFile,
    current_user: Annotated[User, Depends(get_current_active_user)],
):
    is_super_user(current_user)
    with open(f'{TABLES_DIR}/{table.filename}', 'wb') as buffer:
        copyfileobj(table.file, buffer)
    return {'filename': table.filename}


@router.post('/delete_table')
async def delete_table(
    table_names: DeleteTables,
    current_user: Annotated[User, Depends(get_current_active_user)],
):
    is_super_user(current_user)
    map(lambda x: os.remove(f'{TABLES_DIR}/{x}'), table_names.tables_names)
    return {'Detail': 'Success'}


@router.get('/get_tables')
async def get_tables_names(
    current_user: Annotated[User, Depends(get_current_active_user)],
    session: Session = Depends(get_session),
):
    is_super_user(current_user)
    return {'tables': list(map(lambda x: x, os.listdir(TABLES_DIR)))}


@router.post('/add_data_to_base')
async def add_to_base(
    tables: UploadTabelToDb,
    current_user: Annotated[User, Depends(get_current_active_user)],
    session: Session = Depends(get_session),
):
    is_super_user(current_user)
    for name in tables.tables_names:
        all = []
        document = Document(f'{TABLES_DIR}/{name}')
        tablesDocx = document.tables
        paragraphs = document.paragraphs
        address = paragraphs[-2].text
        is_paid = (
            False if paragraphs[0].text.split()[-1] == '(БЮДЖЕТ)' else True
        )
        courses = list(
            map(
                convert_to_models,
                (
                    [k, row, address, is_paid]
                    for k, row in enumerate(tablesDocx[0].rows)
                ),
            )
        )[2:]
        all.extend(courses)
    for course, schedule in all:
        c, s = Course(**course), Schedule(**schedule)
        session.add(s)
        session.commit()
        c.schedule_id = s.id
        session.add(c)
    session.commit()
    return {'Result': 'Success'}


@router.post('/add_course')
async def add_course(
    course: CourseType,
    schedule: ScheduleType,
    current_user: Annotated[User, Depends(get_current_active_user)],
    session: Session = Depends(get_session),
):
    is_super_user(current_user)
    sc = Schedule(**schedule.model_dump())
    session.add(sc)
    session.commit()
    c = course.model_dump()
    c.update({'schedule_id': sc.id})
    session.add(Course(**c))
    session.commit()
    return {'Result': 'Success'}


@router.patch('/change_course/{id}')
async def change_course(
    id: int,
    course: CourseType,
    schedule: ScheduleType,
    current_user: Annotated[User, Depends(get_current_active_user)],
    session: Session = Depends(get_session),
):
    is_super_user(current_user)
    course_model = session.exec(select(Course).where(Course.id == id)).one()
    schedule_model = course_model.schedule
    session.add(schedule_model.sqlmodel_update(schedule.model_dump()))
    session.add(course_model.sqlmodel_update(course.model_dump()))
    session.commit()
    return {'Result': 'Success'}


@router.delete('/delete_course/{id}')
async def delete_course(
    id: int,
    current_user: Annotated[User, Depends(get_current_active_user)],
    session: Session = Depends(get_session),
):
    is_super_user(current_user)
    session.delete(session.exec(select(Course).where(Course.id == id)).one())
    session.delete(
        session.exec(select(Schedule).where(Schedule.id == id)).one()
    )
    session.commit()
    return {'Result': 'Success'}
